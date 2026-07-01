"""
报告导出工具 - 支持 Markdown、Word、PDF 格式

依赖安装:
    pip install pypandoc markdown

PDF 导出需要额外工具:
    - wkhtmltopdf (推荐): https://wkhtmltopdf.org/downloads.html
    - 或 LaTeX: https://www.latex-project.org/get/
"""

import logging
import os
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

# 检查依赖是否可用
try:
    import markdown
    import pypandoc

    # 检查 pandoc 是否可用
    try:
        pypandoc.get_pandoc_version()
        PANDOC_AVAILABLE = True
        logger.info("✅ Pandoc 可用")
    except OSError:
        PANDOC_AVAILABLE = False
        logger.warning("⚠️ Pandoc 不可用，Word 和 PDF 导出功能将不可用")

    EXPORT_AVAILABLE = True
except ImportError as e:
    EXPORT_AVAILABLE = False
    PANDOC_AVAILABLE = False
    logger.warning(f"⚠️ 导出功能依赖包缺失: {e}")
    logger.info("💡 请安装: pip install pypandoc markdown")

# 检查 pdfkit（唯一的 PDF 生成工具）
PDFKIT_AVAILABLE = False
PDFKIT_ERROR = None

try:
    import pdfkit
    # 检查 wkhtmltopdf 是否安装
    try:
        pdfkit.configuration()
        PDFKIT_AVAILABLE = True
        logger.info("✅ pdfkit + wkhtmltopdf 可用（PDF 生成工具）")
    except Exception as e:
        PDFKIT_ERROR = str(e)
        logger.warning("⚠️ wkhtmltopdf 未安装，PDF 导出功能不可用")
        logger.info("💡 安装方法: https://wkhtmltopdf.org/downloads.html")
except ImportError:
    logger.warning("⚠️ pdfkit 未安装，PDF 导出功能不可用")
    logger.info("💡 安装方法: pip install pdfkit")
except Exception as e:
    PDFKIT_ERROR = str(e)
    logger.warning(f"⚠️ pdfkit 检测失败: {e}")


class ReportExporter:
    """报告导出器 - 支持 Markdown、Word、PDF 格式"""

    # 统一的英文 -> 中文模块标题映射
    MODULE_TITLE_MAP = {
        # 报告概览 / 决策 / 摘要
        "report_summary": "📊 报告摘要",
        "summary": "📋 执行摘要",
        "executive_summary": "📌 执行摘要",
        "analysis_summary": "📊 分析摘要",
        "overview": "🔎 综合概览",

        # 分析师团队 (7个)
        "market_report": "📈 市场技术分析",
        "market_analysis": "📈 市场技术分析",
        "sentiment_report": "💭 市场情绪分析",
        "news_report": "📰 新闻事件分析",
        "fundamentals_report": "💰 基本面分析",
        "policy_report": "🏛️ 政策分析",
        "hot_money_report": "💹 游资追踪分析",
        "lockup_report": "🔒 限售解禁分析",

        # 研究团队 (3个)
        "bull_researcher": "🐂 看涨研究员",
        "bear_researcher": "🐻 看跌研究员",
        "research_team_decision": "👔 研究经理决策",
        "investment_debate_state": "🧠 多空辩论",
        "bull_history": "🐂 多头观点",
        "bear_history": "🐻 空头观点",
        "debate_history": "🧐 辩论记录",

        # 交易团队 (1个)
        "trader_investment_plan": "💼 交易员投资计划",
        "trader_investment_decision": "💼 交易员投资决策",

        # 风险管理团队 (5个)
        "risky_analyst": "🔥 激进风险分析",
        "safe_analyst": "🛡️ 保守风险分析",
        "neutral_analyst": "⚖️ 中性风险分析",
        "risk_control_decision": "📋 风控约束决策",
        "risk_management_decision": "👔 风险经理决策",
        "risk_debate_state": "🛡️ 风险辩论",
        "aggressive_history": "🔥 激进观点",
        "conservative_history": "❄️ 保守观点",
        "neutral_history": "⚖️ 中性观点",

        # 最终决策
        "final_trade_decision": "🎯 决策建议",
        "final_decision": "🎯 决策建议",
        "decision_summary": "📌 决策摘要",

        # 数据质量
        "data_quality_summary": "📊 数据质量评估",
        "quality_gate": "🚦 数据质量门控",

        # 市场/宏观
        "macro_report": "📈 宏观分析",
        "market_overview": "🌍 市场概况",

        # 公司/财务
        "company_overview": "🏢 公司概况",
        "financial_analysis": "💰 财务分析",
        "valuation_analysis": "💎 估值分析",

        # 技术面
        "technical_analysis": "📈 技术分析",
        "technical_analysis_report": "📈 技术分析",
        "technical_report": "📈 技术分析",

        # 交易量
        "trading_volume": "📊 交易量分析",
        "volume_report": "📊 成交量分析",

        # 风险
        "risk_analysis": "⚠️ 风险分析",
        "risk_report": "⚠️ 风险提示",
        "risk_level": "⚠️ 风险等级",
        "confidence_score": "🎯 置信度",

        # 投资建议
        "investment_plan": "📝 投资计划",
        "investment_recommendation": "🎯 投资建议",
        "recommendation": "🎯 投资建议",
        "trading_signal": "🚦 交易信号",

        # 行业/板块
        "industry_report": "🏭 行业分析",
        "sector_analysis": "🏭 板块分析",

        # 其他
        "detailed_analysis": "📂 详细分析",
        "state": "📊 完整分析状态",

        # 结构化字段（决策部分）
        "action": "操作建议",
        "target_price": "目标价",
        "stop_loss": "止损价",
        "confidence": "置信度",
        "reasoning": "推理说明",
        "bull_points": "看多理由",
        "bear_points": "看空理由",
        "key_points": "核心要点",
        "technical_signals": "技术信号",
        "judge_decision": "裁判裁决",
        "current_response": "当前回应",
    }

    def get_module_title(self, module_key: str) -> str:
        """获取模块的中文标题，未命中时返回原 key"""
        if module_key in self.MODULE_TITLE_MAP:
            return self.MODULE_TITLE_MAP[module_key]
        # 模糊匹配：去除常见后缀再查
        cleaned = module_key.replace("_report", "").replace("_analysis", "").replace("_state", "")
        if cleaned in self.MODULE_TITLE_MAP:
            return self.MODULE_TITLE_MAP[cleaned]
        return module_key

    def __init__(self):
        self.export_available = EXPORT_AVAILABLE
        self.pandoc_available = PANDOC_AVAILABLE
        self.pdfkit_available = PDFKIT_AVAILABLE

        logger.info("📋 ReportExporter 初始化:")
        logger.info(f"  - export_available: {self.export_available}")
        logger.info(f"  - pandoc_available: {self.pandoc_available}")
        logger.info(f"  - pdfkit_available: {self.pdfkit_available}")

    def generate_markdown_report(self, report_doc: Dict[str, Any]) -> str:
        """生成 Markdown 格式报告（全中文标题）"""
        logger.info("📝 生成 Markdown 报告...")

        stock_symbol = report_doc.get("stock_symbol", "未知")
        analysis_date = report_doc.get("analysis_date", "")
        analysts = report_doc.get("analysts", [])
        reports = report_doc.get("reports", {})
        summary = report_doc.get("summary", "")

        # 从 reports 或顶层提取额外内容（通常由 decision 提供）
        decision = report_doc.get("decision") or report_doc.get("detailed_analysis") or {}

        content_parts = []

        # 标题和元信息
        content_parts.append(f"# {stock_symbol} 股票分析报告")
        content_parts.append("")
        content_parts.append(f"**📅 分析日期**: {analysis_date}")
        if analysts:
            content_parts.append(f"**👤 分析师团队**: {', '.join(analysts)}")
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")

        # 执行摘要
        if summary:
            content_parts.append("## 📊 执行摘要")
            content_parts.append("")
            content_parts.append(summary)
            content_parts.append("")
            content_parts.append("---")
            content_parts.append("")

        # 决策摘要（如果有）
        if decision and isinstance(decision, dict):
            decision_title = decision.get("title") or decision.get("decision_title")
            if decision_title:
                content_parts.append("## 🎯 投资决策")
                content_parts.append("")
                content_parts.append(str(decision_title))
                content_parts.append("")
            action = decision.get("action") or decision.get("recommendation")
            if action:
                content_parts.append(f"- **操作建议**: {action}")
            target_price = decision.get("target_price")
            if target_price:
                content_parts.append(f"- **🎯 目标价**: {target_price}")
            stop_loss = decision.get("stop_loss")
            if stop_loss:
                content_parts.append(f"- **🛑 止损价**: {stop_loss}")
            confidence = decision.get("confidence") or decision.get("confidence_score")
            if confidence:
                content_parts.append(f"- **🔒 置信度**: {confidence}")
            risk = decision.get("risk_level")
            if risk:
                content_parts.append(f"- **⚠️ 风险等级**: {risk}")
            if decision.get("executive_summary"):
                content_parts.append(f"- **📝 决策摘要**: {decision['executive_summary']}")
            if any([decision_title, action, target_price, confidence]):
                content_parts.append("")
                content_parts.append("---")
                content_parts.append("")

        # 各模块内容 - 按推荐顺序遍历（如果 reports 中有这些模块）
        preferred_order = [
            "company_overview",
            "market_report",
            "fundamentals_report",
            "financial_analysis",
            "technical_analysis_report",
            "technical_analysis",
            "sentiment_report",
            "news_report",
            "policy_report",
            "hot_money_report",
            "lockup_report",
            "investment_plan",
            "trader_investment_plan",
            "trader_investment_decision",
            "valuation_analysis",
            "market_analysis",
            "risk_analysis",
            "risk_report",
            "investment_recommendation",
            "final_trade_decision",
            "report_summary",
            "summary",
            "overview",
            "state",
            "detailed_analysis",
        ]

        # 先按推荐顺序输出
        for module_key in preferred_order:
            if module_key in reports:
                module_content = reports[module_key]
                self._append_module_content(content_parts, module_key, module_content)

        # 再输出未包含在优先顺序中的模块（保持原有顺序）
        for module_key, module_content in reports.items():
            if module_key not in preferred_order:
                self._append_module_content(content_parts, module_key, module_content)

        # 页脚
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        content_parts.append("*本报告由 AI 自动生成，仅供研究参考，不构成任何投资建议。*")
        content_parts.append("")

        markdown_content = "\n".join(content_parts)
        logger.info(f"✅ Markdown 报告生成完成，长度: {len(markdown_content)} 字符")

        return markdown_content

    def _append_module_content(
        self,
        content_parts: list,
        module_key: str,
        module_content: Any,
    ) -> None:
        """将一个模块的内容添加到 content_parts"""
        if module_content is None:
            return
        # 字符串：直接写入
        if isinstance(module_content, str):
            if not module_content.strip():
                return
            title = self.get_module_title(module_key)
            content_parts.append(f"## {title}")
            content_parts.append("")
            content_parts.append(module_content)
            content_parts.append("")
            content_parts.append("---")
            content_parts.append("")
            return
        # 字典：递归展开
        if isinstance(module_content, dict):
            title = self.get_module_title(module_key)
            # 如果字典有核心内容（text/result/content），直接取
            for text_key in ("content", "text", "markdown", "output", "result"):
                if text_key in module_content and isinstance(module_content[text_key], str):
                    content_parts.append(f"## {title}")
                    content_parts.append("")
                    content_parts.append(module_content[text_key])
                    content_parts.append("")
                    content_parts.append("---")
                    content_parts.append("")
                    return
            # 否则：展开子键
            content_parts.append(f"## {title}")
            content_parts.append("")
            for sub_key, sub_val in module_content.items():
                if sub_key in ("analysis_id", "stock_symbol", "trade_date", "created_at"):
                    continue
                sub_title = self.get_module_title(sub_key)
                if isinstance(sub_val, str):
                    content_parts.append(f"### {sub_title}")
                    content_parts.append("")
                    content_parts.append(sub_val)
                    content_parts.append("")
                elif isinstance(sub_val, (int, float, bool)):
                    content_parts.append(f"- **{sub_title}**: {sub_val}")
                elif sub_val:
                    content_parts.append(f"### {sub_title}")
                    content_parts.append("")
                    content_parts.append(f"```json\n{sub_val}\n```")
                    content_parts.append("")
            content_parts.append("---")
            content_parts.append("")
        # 列表：简单渲染
        elif isinstance(module_content, list):
            title = self.get_module_title(module_key)
            content_parts.append(f"## {title}")
            content_parts.append("")
            for item in module_content:
                if isinstance(item, str):
                    content_parts.append(f"- {item}")
                else:
                    content_parts.append(f"- {item}")
            content_parts.append("")
            content_parts.append("---")
            content_parts.append("")
    
    def _clean_markdown_for_pandoc(self, md_content: str) -> str:
        """清理 Markdown 内容，避免 pandoc 解析问题"""
        import re

        # 移除可能导致 YAML 解析问题的内容
        # 如果开头有 "---"，在前面添加空行
        if md_content.strip().startswith("---"):
            md_content = "\n" + md_content

        # 🔥 移除可能导致竖排的 HTML 标签和样式
        # 移除 writing-mode 相关的样式
        md_content = re.sub(r'<[^>]*writing-mode[^>]*>', '', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<[^>]*text-orientation[^>]*>', '', md_content, flags=re.IGNORECASE)

        # 移除 <div> 标签中的 style 属性（可能包含竖排样式）
        md_content = re.sub(r'<div\s+style="[^"]*">', '<div>', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<span\s+style="[^"]*">', '<span>', md_content, flags=re.IGNORECASE)

        # 🔥 移除可能导致问题的 HTML 标签
        # 保留基本的 Markdown 格式，移除复杂的 HTML
        md_content = re.sub(r'<style[^>]*>.*?</style>', '', md_content, flags=re.DOTALL | re.IGNORECASE)

        # 🔥 确保所有段落都是正常的横排文本
        # 在每个段落前后添加明确的换行，避免 Pandoc 误判
        lines = md_content.split('\n')
        cleaned_lines = []
        for line in lines:
            # 跳过空行
            if not line.strip():
                cleaned_lines.append(line)
                continue

            # 如果是标题、列表、表格等 Markdown 语法，保持原样
            if line.strip().startswith(('#', '-', '*', '|', '>', '```', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                cleaned_lines.append(line)
            else:
                # 普通段落：确保没有特殊字符导致竖排
                cleaned_lines.append(line)

        md_content = '\n'.join(cleaned_lines)

        return md_content

    def _create_pdf_css(self) -> str:
        """创建 PDF 样式表，控制表格分页和文本方向"""
        return """
<style>
/* 🔥 强制所有文本横排显示（修复中文竖排问题） */
* {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

body {
    writing-mode: horizontal-tb !important;
    direction: ltr !important;
}

/* 段落和文本 */
p, div, span, td, th, li {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* 表格样式 - 允许跨页 */
table {
    width: 100%;
    border-collapse: collapse;
    page-break-inside: auto;
    writing-mode: horizontal-tb !important;
}

/* 表格行 - 避免在行中间分页 */
tr {
    page-break-inside: avoid;
    page-break-after: auto;
}

/* 表头 - 在每页重复显示 */
thead {
    display: table-header-group;
}

/* 表格单元格 */
td, th {
    padding: 8px;
    border: 1px solid #ddd;
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* 表头样式 */
th {
    background-color: #f2f2f2;
    font-weight: bold;
}

/* 避免标题后立即分页 */
h1, h2, h3, h4, h5, h6 {
    page-break-after: avoid;
    writing-mode: horizontal-tb !important;
}

/* 避免在列表项中间分页 */
li {
    page-break-inside: avoid;
}

/* 代码块 */
pre, code {
    writing-mode: horizontal-tb !important;
    white-space: pre-wrap;
    word-wrap: break-word;
}
</style>
"""
    
    def generate_docx_report(self, report_doc: Dict[str, Any]) -> bytes:
        """生成 Word 文档格式报告"""
        logger.info("📄 开始生成 Word 文档...")

        if not self.pandoc_available:
            raise Exception("Pandoc 不可用，无法生成 Word 文档。请安装 pandoc 或使用 Markdown 格式导出。")

        # 生成 Markdown 内容
        md_content = self.generate_markdown_report(report_doc)

        try:
            # 创建临时文件（使用纯ASCII路径避免编码问题）
            import os as os_module
            tmp_dir = tempfile.gettempdir()
            # 使用纯ASCII文件名
            output_file = os.path.join(tmp_dir, f"report_{os_module.getpid()}_{id(self) % 10000}.docx")
            
            logger.info(f"📁 临时文件路径: {output_file}")

            # Pandoc 参数
            extra_args = [
                '--from=markdown-yaml_metadata_block',  # 禁用 YAML 元数据块解析
                '--standalone',  # 生成独立文档
                '--wrap=preserve',  # 保留换行
                '--columns=120',  # 设置列宽
                '-M', 'lang=zh-CN',  # 🔥 明确指定语言为简体中文
                '-M', 'dir=ltr',  # 🔥 明确指定文本方向为从左到右
            ]

            # 清理内容
            cleaned_content = self._clean_markdown_for_pandoc(md_content)
            # 确保内容是有效的UTF-8字符串
            if isinstance(cleaned_content, bytes):
                cleaned_content = cleaned_content.decode('utf-8', errors='replace')
            # 规范化内容，移除可能引起编码问题的特殊字符
            cleaned_content = cleaned_content.encode('utf-8', errors='replace').decode('utf-8', errors='replace')

            # 转换为 Word
            pypandoc.convert_text(
                cleaned_content,
                'docx',
                format='markdown',
                outputfile=output_file,
                extra_args=extra_args
            )

            logger.info("✅ pypandoc 转换完成")

            # 🔥 后处理：修复 Word 文档中的文本方向
            try:
                from docx import Document
                doc = Document(output_file)

                # 修复所有段落的文本方向
                for paragraph in doc.paragraphs:
                    # 设置段落为从左到右
                    if paragraph._element.pPr is not None:
                        # 移除可能的竖排设置
                        for child in list(paragraph._element.pPr):
                            if 'textDirection' in child.tag or 'bidi' in child.tag:
                                paragraph._element.pPr.remove(child)

                # 修复表格中的文本方向
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph._element.pPr is not None:
                                    for child in list(paragraph._element.pPr):
                                        if 'textDirection' in child.tag or 'bidi' in child.tag:
                                            paragraph._element.pPr.remove(child)

                # 保存修复后的文档
                doc.save(output_file)
                logger.info("✅ Word 文档文本方向修复完成")
            except ImportError:
                logger.warning("⚠️ python-docx 未安装，跳过文本方向修复")
            except Exception as e:
                logger.warning(f"⚠️ Word 文档文本方向修复失败: {e}")

            # 读取生成的文件
            with open(output_file, 'rb') as f:
                docx_content = f.read()

            logger.info(f"✅ Word 文档生成成功，大小: {len(docx_content)} 字节")

            # 清理临时文件
            os.unlink(output_file)

            return docx_content
            
        except Exception as e:
            logger.error(f"❌ Word 文档生成失败: {e}", exc_info=True)
            # 清理临时文件
            try:
                if 'output_file' in locals() and os.path.exists(output_file):
                    os.unlink(output_file)
            except:
                pass
            raise Exception(f"生成 Word 文档失败: {e}")
    
    def _markdown_to_html(self, md_content: str) -> str:
        """将 Markdown 转换为 HTML"""
        import markdown

        # 配置 Markdown 扩展
        extensions = [
            'markdown.extensions.tables',  # 表格支持
            'markdown.extensions.fenced_code',  # 代码块支持
            'markdown.extensions.nl2br',  # 换行支持
        ]

        # 转换为 HTML
        html_content = markdown.markdown(md_content, extensions=extensions)

        # 添加 HTML 模板和样式
        # WeasyPrint 优化的 CSS（移除不支持的属性）
        html_template = f"""
<!DOCTYPE html>
<html lang="zh-CN" dir="ltr">
<head>
    <meta charset="UTF-8">
    <title>分析报告</title>
    <style>
        /* 基础样式 - 确保文本方向正确 */
        html {{
            direction: ltr;
        }}

        body {{
            font-family: "Noto Sans CJK SC", "Microsoft YaHei", "SimHei", "Arial", sans-serif;
            line-height: 1.8;
            color: #333;
            margin: 20mm;
            padding: 0;
            background: white;
            direction: ltr;
        }}

        /* 标题样式 */
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 1.5em;
            margin-bottom: 0.8em;
            font-weight: 600;
            page-break-after: avoid;
            direction: ltr;
        }}

        h1 {{
            font-size: 2em;
            border-bottom: 3px solid #3498db;
            padding-bottom: 0.3em;
            page-break-before: always;
        }}

        h1:first-child {{
            page-break-before: avoid;
        }}

        h2 {{
            font-size: 1.6em;
            border-bottom: 2px solid #bdc3c7;
            padding-bottom: 0.25em;
        }}

        h3 {{
            font-size: 1.3em;
            color: #34495e;
        }}

        /* 段落样式 */
        p {{
            margin: 0.8em 0;
            text-align: left;
            direction: ltr;
        }}

        /* 表格样式 - 优化分页 */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1.5em 0;
            font-size: 0.9em;
            direction: ltr;
        }}

        /* 表头在每页重复 */
        thead {{
            display: table-header-group;
        }}

        tbody {{
            display: table-row-group;
        }}

        /* 表格行避免跨页断开 */
        tr {{
            page-break-inside: avoid;
        }}

        th, td {{
            border: 1px solid #ddd;
            padding: 10px 12px;
            text-align: left;
            direction: ltr;
        }}

        th {{
            background-color: #3498db;
            color: white;
            font-weight: bold;
        }}

        tbody tr:nth-child(even) {{
            background-color: #f8f9fa;
        }}

        tbody tr:hover {{
            background-color: #e9ecef;
        }}

        /* 代码块样式 */
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: "Consolas", "Monaco", "Courier New", monospace;
            font-size: 0.9em;
            direction: ltr;
        }}

        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            border-left: 4px solid #3498db;
            page-break-inside: avoid;
            direction: ltr;
        }}

        pre code {{
            background-color: transparent;
            padding: 0;
        }}

        /* 列表样式 */
        ul, ol {{
            margin: 0.8em 0;
            padding-left: 2em;
            direction: ltr;
        }}

        li {{
            margin: 0.4em 0;
            direction: ltr;
        }}

        /* 强调文本 */
        strong, b {{
            font-weight: 700;
            color: #2c3e50;
        }}

        em, i {{
            font-style: italic;
            color: #555;
        }}

        /* 水平线 */
        hr {{
            border: none;
            border-top: 2px solid #ecf0f1;
            margin: 2em 0;
        }}

        /* 链接样式 */
        a {{
            color: #3498db;
            text-decoration: none;
        }}

        a:hover {{
            text-decoration: underline;
        }}

        /* 分页控制 */
        @page {{
            size: A4;
            margin: 20mm;

            @top-center {{
                content: "分析报告";
                font-size: 10pt;
                color: #999;
            }}

            @bottom-right {{
                content: "第 " counter(page) " 页";
                font-size: 10pt;
                color: #999;
            }}
        }}

        /* 避免孤行和寡行 */
        p, li {{
            orphans: 3;
            widows: 3;
        }}

        /* 图片样式 */
        img {{
            max-width: 100%;
            height: auto;
            page-break-inside: avoid;
        }}

        /* 引用块样式 */
        blockquote {{
            margin: 1em 0;
            padding: 0.5em 1em;
            border-left: 4px solid #3498db;
            background-color: #f8f9fa;
            font-style: italic;
            page-break-inside: avoid;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
        return html_template

    def _generate_pdf_with_pdfkit(self, html_content: str) -> bytes:
        """使用 pdfkit 生成 PDF（先尝试直接字符串方式，失败则使用命令行方式）"""
        import pdfkit
        import tempfile
        import subprocess

        logger.info("🔧 使用 pdfkit + wkhtmltopdf 生成 PDF...")

        # 配置选项
        options = {
            'encoding': 'UTF-8',
            'enable-local-file-access': None,
            'page-size': 'A4',
            'margin-top': '20mm',
            'margin-right': '20mm',
            'margin-bottom': '20mm',
            'margin-left': '20mm',
            'quiet': '',
        }

        # 方法1：尝试直接使用字符串方式
        try:
            # 确保内容是字符串
            if isinstance(html_content, bytes):
                html_content = html_content.decode('utf-8')
            
            pdf_bytes = pdfkit.from_string(html_content, False, options=options)
            logger.info(f"✅ pdfkit PDF 生成成功（字符串方式），大小: {len(pdf_bytes)} 字节")
            return pdf_bytes
        except Exception as e:
            logger.warning(f"⚠️ 字符串方式失败，尝试命令行方式: {e}")

        # 方法2：使用临时文件 + 命令行方式
        try:
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', encoding='utf-8', delete=False) as f:
                f.write(html_content)
                html_file = f.name
            
            with tempfile.NamedTemporaryFile(mode='wb', suffix='.pdf', delete=False) as f:
                pdf_file = f.name
            
            # 构建命令行参数
            cmd = [
                'wkhtmltopdf',
                '--encoding', 'UTF-8',
                '--enable-local-file-access',
                '--page-size', 'A4',
                '--margin-top', '20mm',
                '--margin-right', '20mm',
                '--margin-bottom', '20mm',
                '--margin-left', '20mm',
                '--quiet',
                html_file,
                pdf_file
            ]
            
            # 执行命令
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode != 0:
                error_msg = f"wkhtmltopdf 执行失败: {result.stderr}"
                logger.error(f"❌ {error_msg}")
                raise Exception(error_msg)
            
            # 读取生成的 PDF
            with open(pdf_file, 'rb') as f:
                pdf_bytes = f.read()
            
            # 清理临时文件
            os.unlink(html_file)
            os.unlink(pdf_file)
            
            logger.info(f"✅ wkhtmltopdf PDF 生成成功（命令行方式），大小: {len(pdf_bytes)} 字节")
            return pdf_bytes
            
        except Exception as e:
            # 清理临时文件（如果存在）
            if 'html_file' in locals():
                try:
                    os.unlink(html_file)
                except:
                    pass
            if 'pdf_file' in locals():
                try:
                    os.unlink(pdf_file)
                except:
                    pass
            logger.error(f"❌ PDF 生成失败（命令行方式）: {e}")
            raise

    def _clean_text(self, text: Any) -> str:
        """清理文本，确保可以正确编码"""
        if text is None:
            return ""
        if isinstance(text, str):
            # 移除无法编码的字符和控制字符
            cleaned = text.encode('utf-8', errors='replace').decode('utf-8', errors='replace')
            # 移除控制字符（除了常见的换行符）
            cleaned = ''.join(c for c in cleaned if c == '\n' or c == '\r' or c == '\t' or ord(c) >= 32)
            return cleaned
        # 其他类型转换为字符串
        return str(text)
    
    def _debug_text_encoding(self, text: str, label: str = "text") -> None:
        """调试文本编码问题"""
        problematic_chars = []
        for i, c in enumerate(text):
            ord_c = ord(c)
            # 检查是否是无法用 latin-1 编码的字符
            if ord_c > 255:
                problematic_chars.append(f"位置 {i}: '{c}' (U+{ord_c:04X})")
        
        if problematic_chars:
            logger.warning(f"⚠️ 发现 {len(problematic_chars)} 个无法用 latin-1 编码的字符 ({label}):")
            for pc in problematic_chars[:10]:  # 最多显示10个
                logger.warning(f"  {pc}")
            if len(problematic_chars) > 10:
                logger.warning(f"  ... 还有 {len(problematic_chars) - 10} 个")
    
    def _clean_report_data(self, report_doc: Dict[str, Any]) -> Dict[str, Any]:
        """递归清理报告数据，确保所有字符串都是有效的"""
        cleaned = {}
        for key, value in report_doc.items():
            if isinstance(value, dict):
                cleaned[key] = self._clean_report_data(value)
            elif isinstance(value, list):
                cleaned[key] = [self._clean_text(item) if isinstance(item, (str, bytes)) else item for item in value]
            elif isinstance(value, str):
                cleaned[key] = self._clean_text(value)
            else:
                cleaned[key] = value
        return cleaned
    
    def generate_pdf_report(self, report_doc: Dict[str, Any]) -> bytes:
        """生成 PDF 格式报告（使用 stdin/stdout 传递数据，避免编码问题）"""
        import subprocess
        
        logger.info("📊 开始生成 PDF 文档...")

        # 清理报告数据，确保所有字符串都是有效的 UTF-8
        cleaned_doc = self._clean_report_data(report_doc)
        logger.info("✅ 报告数据清理完成")

        # 生成 Markdown 内容
        md_content = self.generate_markdown_report(cleaned_doc)
        logger.info(f"📝 Markdown 内容长度: {len(md_content)} 字符")
        
        # 调试：检查 Markdown 内容的编码
        self._debug_text_encoding(md_content, "Markdown")

        # 清理 Markdown 内容
        md_content = self._clean_text(md_content)
        logger.info(f"📝 清理后 Markdown 内容长度: {len(md_content)} 字符")

        # 转换为 HTML
        html_content = self._markdown_to_html(md_content)
        logger.info(f"📝 HTML 内容长度: {len(html_content)} 字符")
        
        # 调试：检查 HTML 内容的编码
        self._debug_text_encoding(html_content, "HTML")
        
        # 清理 HTML 内容
        html_content = self._clean_text(html_content)
        logger.info(f"📝 清理后 HTML 内容长度: {len(html_content)} 字符")
        
        # 使用 stdin/stdout 方式传递数据，避免文件编码问题
        try:
            # 构建命令行参数（使用 - 表示从 stdin 读取，输出到 stdout）
            cmd = [
                'wkhtmltopdf',
                '--encoding', 'UTF-8',
                '--enable-local-file-access',
                '--page-size', 'A4',
                '--margin-top', '20mm',
                '--margin-right', '20mm',
                '--margin-bottom', '20mm',
                '--margin-left', '20mm',
                '--quiet',
                '-',  # 从 stdin 读取 HTML
                '-'   # 输出到 stdout
            ]
            
            logger.info(f"🔧 执行命令: {' '.join(cmd)}")
            
            # 执行命令（使用字节模式，直接通过 stdin/stdout 传递数据）
            # 将 HTML 内容编码为 UTF-8 字节
            html_bytes = html_content.encode('utf-8', errors='replace')
            logger.info(f"📤 输入 HTML 字节长度: {len(html_bytes)}")
            
            result = subprocess.run(cmd, input=html_bytes, capture_output=True)
            
            logger.info(f"📊 命令执行返回码: {result.returncode}")
            if result.stdout:
                logger.info(f"📋 标准输出长度: {len(result.stdout)} 字节")
            if result.stderr:
                try:
                    logger.info(f"📋 标准错误: {result.stderr.decode('utf-8', errors='replace')[:200]}")
                except:
                    logger.info(f"📋 标准错误: (二进制数据，长度: {len(result.stderr)})")
            
            if result.returncode != 0:
                try:
                    error_msg = f"wkhtmltopdf 执行失败: {result.stderr.decode('utf-8', errors='replace')}"
                except:
                    error_msg = f"wkhtmltopdf 执行失败，返回码: {result.returncode}"
                logger.error(f"❌ {error_msg}")
                raise Exception(error_msg)
            
            # 如果返回码为0但stdout为空，也认为失败
            if not result.stdout or len(result.stdout) < 100:
                error_msg = "wkhtmltopdf 生成的 PDF 数据为空或太小"
                logger.error(f"❌ {error_msg}")
                raise Exception(error_msg)
            
            # 标准输出就是 PDF 数据
            pdf_bytes = result.stdout
            logger.info(f"✅ PDF 生成成功，大小: {len(pdf_bytes)} 字节")
            return pdf_bytes
            
        except FileNotFoundError:
            error_msg = "wkhtmltopdf 命令未找到，请先安装 wkhtmltopdf"
            logger.error(f"❌ {error_msg}")
            raise Exception(error_msg)
        except Exception as e:
            error_msg = f"PDF 生成失败: {e}"
            logger.error(f"❌ {error_msg}")
            raise Exception(error_msg)
    
    def _generate_pdf_with_pandoc(self, md_content: str) -> bytes:
        """使用 pandoc 生成 PDF（备选方案）"""
        import subprocess
        
        # 构建命令行参数
        cmd = [
            'pandoc',
            '-f', 'markdown',
            '-t', 'pdf',
            '--pdf-engine=xelatex',
            '-V', 'mainfont=SimSun',
            '-V', 'sansfont=SimHei',
            '-V', 'CJKmainfont=SimSun',
            '-o', '-'  # 输出到 stdout
        ]
        
        logger.info(f"🔧 执行 pandoc 命令: {' '.join(cmd)}")
        
        # 将 Markdown 内容编码为 UTF-8
        md_bytes = md_content.encode('utf-8', errors='replace')
        
        result = subprocess.run(cmd, input=md_bytes, capture_output=True)
        
        if result.returncode != 0:
            try:
                error_msg = f"pandoc 执行失败: {result.stderr.decode('utf-8', errors='replace')}"
            except:
                error_msg = f"pandoc 执行失败，返回码: {result.returncode}"
            logger.error(f"❌ {error_msg}")
            raise Exception(error_msg)
        
        return result.stdout


# 创建全局导出器实例
report_exporter = ReportExporter()

